import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE

doc = Document()

# Adjust margins to match Tushar's resume layout (0.5" left/right, 0.6" top/bottom)
for section in doc.sections:
    section.top_margin    = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin   = Inches(0.5)
    section.right_margin  = Inches(0.5)

FONT = 'Times New Roman'

# Helper function to add a grey horizontal rule under section headers
def add_hrule(para):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')  # Slightly thinner thickness
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '7F7F7F')  # Professional grey divider
    pBdr.append(bottom)
    pPr.append(pBdr)

# Helper to add hyperlinks in python-docx
def add_hyperlink(paragraph, url, text, color="000000", underline=True, size_pt=9.0):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT)
    rFonts.set(qn('w:hAnsi'), FONT)
    rPr.append(rFonts)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size_pt * 2)))  # half-points
    rPr.append(sz)

    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)

    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

# Helper to add multiple right-aligned items (such as links separated by |)
def add_right_items(paragraph, items, font_size=9.2, italic=False):
    # Add tab character to push to the right-aligned tab stop
    paragraph.add_run("\t")
    
    if not isinstance(items, list):
        items = [items]
        
    first = True
    for item in items:
        if not first:
            r_sep = paragraph.add_run(" | ")
            r_sep.font.name = FONT
            r_sep.font.size = Pt(font_size)
            r_sep.font.color.rgb = RGBColor(0, 0, 0)
            r_sep.italic = italic
            
        first = False
        
        if isinstance(item, tuple):
            text, url = item
            if url:
                add_hyperlink(paragraph, url, text, color="000000", underline=True, size_pt=font_size)
            else:
                r = paragraph.add_run(text)
                r.font.name = FONT
                r.font.size = Pt(font_size)
                r.font.color.rgb = RGBColor(0, 0, 0)
                r.italic = italic
        else:
            r = paragraph.add_run(item)
            r.font.name = FONT
            r.font.size = Pt(font_size)
            r.font.color.rgb = RGBColor(0, 0, 0)
            r.italic = italic

# Section header styling
def section_header(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = FONT
    run.font.color.rgb = RGBColor(0, 0, 0)  # Pure black
    add_hrule(p)
    return p

# Standard One-Line Header (for Projects)
def one_line_header(left_text, right_items, font_size=9.5, space_before=4, space_after=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
    
    r = p.add_run(left_text)
    r.font.name = FONT
    r.font.size = Pt(font_size)
    r.font.color.rgb = RGBColor(0, 0, 0)
    r.bold = True
    
    add_right_items(p, right_items, font_size=font_size)
    return p

# Standard Two-Line Header (for Experience, Education, Certifications, Leadership)
def two_line_header(left_l1, right_l1, left_l2, right_l2, font_size=9.5, space_before=4, space_after=1):
    # Line 1: Bold title left, Date/Links right
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(space_before)
    p1.paragraph_format.space_after  = Pt(0)
    p1.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
    
    r1 = p1.add_run(left_l1)
    r1.font.name = FONT
    r1.font.size = Pt(font_size)
    r1.font.color.rgb = RGBColor(0, 0, 0)
    r1.bold = True
    
    add_right_items(p1, right_l1, font_size=font_size)
    
    # Line 2: Italic Organization left, Italic Location/Details right
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(space_after)
    p2.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
    
    r2 = p2.add_run(left_l2)
    r2.font.name = FONT
    r2.font.size = Pt(font_size - 0.5)
    r2.font.color.rgb = RGBColor(0, 0, 0)
    r2.italic = True
    
    add_right_items(p2, right_l2, font_size=font_size - 0.5, italic=True)
    
    return p1, p2

# Bullet point styling with bold text parser (**text**)
def bullet(text, size=9.0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1.5)
    p.paragraph_format.left_indent  = Inches(0.2)
    
    parts = text.split("**")
    is_bold = False
    for part in parts:
        run = p.add_run(part)
        run.font.size = Pt(size)
        run.font.name = FONT
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = is_bold
        is_bold = not is_bold

# Plain text styling with bold parser
def plain(text, bold=False, size=9.0, center=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    parts = text.split("**")
    is_bold_part = bold
    for part in parts:
        run = p.add_run(part)
        run.font.size = Pt(size)
        run.font.name = FONT
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = is_bold_part if not bold else True
        run.italic = italic
        is_bold_part = not is_bold_part
    return p


# ── NAME & CONTACT ────────────────────────────────────────────────────────────
p_name = doc.add_paragraph()
p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_name.paragraph_format.space_before = Pt(0)
p_name.paragraph_format.space_after  = Pt(1)
r_name = p_name.add_run("DEVIKA YANAMALA")
r_name.bold = True
r_name.font.size = Pt(15)
r_name.font.name = FONT
r_name.font.color.rgb = RGBColor(0, 0, 0)  # Pure Black

p_contact = doc.add_paragraph()
p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_contact.paragraph_format.space_after = Pt(6)

r_contact = p_contact.add_run("Newcastle, UK  |  +44 7352683914  |  devikayanamala55@gmail.com  |  ")
r_contact.font.size = Pt(9.0)
r_contact.font.name = FONT
r_contact.font.color.rgb = RGBColor(0, 0, 0)

# Clickable, hidden URL links
add_hyperlink(p_contact, "https://www.linkedin.com/in/devika-yanamala-325250345", "LinkedIn", color="000000", underline=True, size_pt=9.0)
r_sep = p_contact.add_run("  |  ")
r_sep.font.size = Pt(9.0)
r_sep.font.name = FONT
r_sep.font.color.rgb = RGBColor(0, 0, 0)
add_hyperlink(p_contact, "https://github.com/DevikaYanamala", "GitHub", color="000000", underline=True, size_pt=9.0)


# ── PROFESSIONAL SUMMARY ──────────────────────────────────────────────────────
section_header("Professional Summary")
plain(
    "Data Science & AI professional with nearly 3 years of hands-on experience in data analytics using "
    "Python, SQL, and Power BI. Experienced in applying machine learning and AI-driven techniques to "
    "analyse data, build end-to-end data pipelines, and generate actionable insights in high-stakes environments. "
    "Skilled at translating complex datasets into clear, commercial stories and business recommendations for stakeholders."
)


# ── TECHNICAL EXPERTISE ───────────────────────────────────────────────────────
section_header("Technical Expertise")
bullet("**Analytics & Tools**: Excel (XLOOKUP, PivotTables, Power Query), Power BI (DAX, Star Schema, interactive dashboards), KPI Reporting & Dashboards, Tableau, Matplotlib, Seaborn, Exploratory Data Analysis (EDA).")
bullet("**Data & Automation**: SQL (PostgreSQL, Snowflake), Python, ETL Data Pipelines, Data Validation & Auditing, exposure to enterprise data systems and SAP-based workflows.")
bullet("**Operations Analytics**: Capacity & Resource Planning, Production Scheduling, Inventory Tracking, Data Collection, Bottleneck Identification, operational efficiency metrics.")
bullet("**AI & Data Science**: Machine Learning (classification, regression, clustering), Predictive Modeling, XGBoost, Scikit-Learn, NLP (basic), exposure to generative AI/LLM tools and workflows, Statistical Analysis (hypothesis testing, time series, A/B testing).")


# ── PROFESSIONAL EXPERIENCE ───────────────────────────────────────────────────
section_header("Professional Experience")

two_line_header(
    left_l1="Global Industry Project – Industry Consulting Experience",
    right_l1="June 2026 – July 2026",
    left_l2="Newcastle University",
    right_l2="Industry Partner Project - Newcastle, UK"
)
bullet("Selected for a competitive, high-impact industry consulting program solving complex operational challenges for a UK industry partner.")
bullet("Collaborated in a team of data scientists to interrogate and analyze historical operational data, identifying critical process bottlenecks.")
bullet("Presented findings and recommendations to stakeholders, focusing on problem-solving and business impact.")

two_line_header(
    left_l1="Data Analyst – Product Analytics",
    right_l1="04/2026 – Present",
    left_l2="Rooted Platform",
    right_l2="Independent Project - Remote"
)
bullet("Engineered a data analytics MVP using Python and SQL to ingest and analyze user behavior trends, user traffic patterns, and authentication logs, enabling data-driven product decisions.")
bullet("Designed and automated interactive Power BI/Streamlit reporting dashboards with advanced DAX measures, improving reporting efficiency and decision-making speed.")
bullet("Applied exploratory data analysis (EDA) and user segmentation to optimize pricing strategy, location-based store recommendations, and search features.")

two_line_header(
    left_l1="Data Analyst (Studentship)",
    right_l1="01/2026 – Present",
    left_l2="Newcastle University",
    right_l2="Part-Time Role - Newcastle upon Tyne, UK"
)
bullet("Built robust Python ETL pipelines for automated data collection, cleaning, validation, and transformation, ensuring 100% data integrity for library services and occupancy analysis.")
bullet("Developed comprehensive KPI dashboards using Power BI to monitor real-time department occupancy, supporting strategic operational decisions.")
bullet("Interrogated large-scale occupancy datasets using SQL and Python to identify spatial usage trends, directly supporting capacity planning and optimizing department resource allocation.")

two_line_header(
    left_l1="Associate Data Scientist",
    right_l1="05/2023 – 05/2025",
    left_l2="BEPEC Solutions",
    right_l2="Full-Time - Bangalore, India"
)
bullet("Developed, tuned, and deployed production-grade machine learning models (classification, regression) in Python, improving overall operational efficiency by 30% and saving 10+ hours per week.")
bullet("Analyzed large-scale structured datasets to identify commercial trends, performance inefficiencies, and logistics bottlenecks, translating complex results into clear business requirements.")
bullet("Partnered with product managers and data engineers in Agile sprints, leading client-facing meetings to gather analytics scope and deliver technical findings.")

two_line_header(
    left_l1="Data Scientist Intern",
    right_l1="01/2023 – 05/2023",
    left_l2="BEPEC Solutions",
    right_l2="Internship - Bangalore, India"
)
bullet("Conducted extensive exploratory data analysis (EDA) and data preprocessing (imputation, encoding, scaling), contributing to a 15% improvement in forecasting accuracy by resolving data quality issues upstream.")
bullet("Created dynamic reporting dashboards in Power BI and Excel to track historical performance metrics, supporting resource planning and department scheduling.")


# ── SELECTED PROJECTS & INITIATIVES ───────────────────────────────────────────
section_header("Selected Projects & Initiatives")

one_line_header(
    left_text="Commercial Opportunities Analytics Dashboard (Retail)",
    right_items=[("Personal Project", None), ("GitHub", "https://github.com/DevikaYanamala/End-to-End-Data-analytics-project-Dmart-")]
)
bullet("Built a Star Schema data model using Python (Pandas) to clean, transform, and optimize raw transactional sales data across 15 European markets.")
bullet("Developed an executive Power BI dashboard with 20+ advanced DAX measures (time intelligence, segment grouping) to track key retail metrics.")
bullet("Analyzed pricing and customer segments to identify discount optimization opportunities, demonstrating over £150K+ in annual profit recovery.")

one_line_header(
    left_text="Pharmaceutical Marketing Strategy & Attribution Analysis",
    right_items=[("Personal Project", None), ("GitHub", "https://github.com/DevikaYanamala/Drug-marketing-compaign-strategy-DA-")]
)
bullet("Built a regression-based marketing attribution model in Python to estimate the commercial impact and ROI of three distinct campaign strategies.")
bullet("Applied Interrupted Time Series (ITS) analysis and statistical testing to evaluate competitor market entry and predict sales deceleration.")
bullet("Designed an interactive Power BI Strategy Simulator using DAX what-if parameters, identifying high-impact channels and driving $1.3B+ in attributed revenue.")

one_line_header(
    left_text="Next Best Action & Churn Prediction Model (Banking)",
    right_items=[("Personal Project", None), ("GitHub", "https://github.com/DevikaYanamala/Data-Scientist-Project")]
)
bullet("Developed an end-to-end machine learning pipeline using Python (Scikit-Learn, Imbalanced-Learn) to predict customer churn, resolving class imbalance with SMOTE.")
bullet("Engineered a Next Best Action (NBA) recommendation algorithm that maps customer transaction patterns to tailored retention products, reducing manual review overhead.")
bullet("Applied SHAP (SHapley Additive exPlanations) to ensure model transparency and explainability, aligning with UK banking regulatory standards.")


# ── EDUCATION ─────────────────────────────────────────────────────────────────
section_header("Education")

two_line_header(
    left_l1="Master of Science (MSc), Data Science & Artificial Intelligence",
    right_l1="Expected 09/2026",
    left_l2="Newcastle University",
    right_l2="Newcastle upon Tyne, UK"
)
bullet("**UK Coursework**: Data Visualisation, Machine Learning, Statistical Foundations of Data Science, Deep Learning, Image Processing, Generative AI for Business, Advanced AI, Group Project.")
bullet("**Dissertation**: Towards the Modelling of Traffic Flows from Open-Source Data – An Assessment of Automated OSM Topological Verification for Machine Learning-Based AADT Prediction. Developing a python pipeline (OSMnx, NetworkX) to resolve topological errors and engineer graph-theoretic centrality features to improve Random Forest and XGBoost traffic prediction models.")

two_line_header(
    left_l1="Bachelor of Technology",
    right_l1="2019 – 2023",
    left_l2="IIIT, RGUKT",
    right_l2="CGPA: 8.7/10 - RK Valley, India"
)
bullet("Relevant Coursework: Data Structures, Algorithms, Database Management Systems, Machine Learning Foundations.")


# ── PROFESSIONAL CERTIFICATIONS ───────────────────────────────────────────────
section_header("Professional Certifications")

two_line_header(
    left_l1="UK Work Culture Programme (Awarded Participant)",
    right_l1="Awarded 2026",
    left_l2="Newcastle University",
    right_l2="Professional Development Programme"
)
bullet("Selected for a competitive programme focused on UK industry standards and corporate communication, presenting business solutions to stakeholders.")


# ── ADDITIONAL EXPERIENCE & RECOGNITION ───────────────────────────────────────
section_header("Additional Experience & Recognition")

two_line_header(
    left_l1="Global Ambassador",
    right_l1="05/2023 – Present",
    left_l2="WomenTech Network",
    right_l2="Volunteering Leadership Role"
)
bullet("Active participant in global initiatives promoting technical excellence, leadership, and diversity in AI and Data Science.")

two_line_header(
    left_l1="Volunteer Module Representative",
    right_l1="01/2026 – Present",
    left_l2="Newcastle University",
    right_l2="Elected Student Leadership Role"
)
bullet("Elected to represent the MSc cohort, dedicating time to collaborate directly with faculty to optimize curriculum delivery and the student experience.")


# Save files to both folders
doc.save(r'c:\Users\devik\Downloads\Devika_Yanamala_CV.docx')
doc.save(r'c:\Users\devik\Downloads\Data-Scientist-Project\Portfolio\Devika_Yanamala_CV.docx')
print("Devika's CV successfully aligned with the senior's resume structure.")
